import os
import json
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document
import time

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API Key Issue")

client=Groq(api_key=my_api_key)
model="llama-3.3-70b-versatile"

#Step 1: Extracting required fields of Job Description in json format
class JobDesc(BaseModel):
    role:str
    required_skills:list[str]
    preffered_skills:list[str]
    min_exp:float | None
    edu_req:list[str]
    responsibilities:list[str]

JobDesc_schema=JobDesc.model_json_schema()
response_format={
    "type":"json_object"
}

system_prompt=f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{JobDesc_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.



If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.

"""
job_description="""
Description
Do you want to solve real customer problems through innovative technology? Do you enjoy working on scalable services in a collaborative team environment? Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. Customer obsession is part of our company DNA, which has made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve complex problems while seeing their work's impact first-hand. The challenges SDEs solve at Amazon are meaningful and influence millions of customers, sellers, and products globally. We seek individuals passionate about creating new products, features, and services while managing ambiguity in an environment where development cycles are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own the entire lifecycle of your code - from design through deployment and ongoing operations. This ownership mindset, combined with our commitment to operational excellence, ensures we deliver the highest quality solutions for our customers.

We're looking for curious minds who think big and want to define tomorrow's technology. At Amazon, you'll grow into the high-impact engineer you know you can be, supported by a culture of learning and mentorship. Every day brings exciting new challenges and opportunities for personal growth.
Key job responsibilities
• Collaborate and communicate effectively with experienced cross-disciplinary Amazonians to design, build, and operate innovative products and services that delight our customers, while participating in technical discussions to drive solutions forward.
• Design and develop scalable solutions using cloud-native architectures and microservices in a large distributed computing environment.
• Participate in code reviews and contribute to technical documentation.
• Build and maintain resilient distributed systems that are scalable, fault-tolerant, and cost-effective.
• Leverage and contribute to the development of GenAI and AI-powered tools to enhance development productivity while staying current with emerging technologies.
• Write clean, maintainable code following best practices and design patterns.
• Work in an agile environment practicing CI/CD principles while participating in operational responsibilities including on-call duties.
• Demonstrate operational excellence through monitoring, troubleshooting, and resolving production issues.
Basic Qualifications
- Experience with at least one general-purpose programming language such as Java, Python, C++, C#, Go, Rust, or TypeScript
- Experience with data structure implementation, basic algorithm development, and/or object-oriented design principles
- Currently has, or is in the process of obtaining a bachelor’s degree in Computer Science, Computer Engineering, Data Science, Information Systems, or related STEM fields
- Must be 18 years of age of older
Preferred Qualifications
- Experience from previous technical internship(s) or demonstrated project experience
- Experience with one or more of the following: AI tools for development productivity, Cloud platforms (preferably AWS), Database systems (SQL and NoSQL), Contributing to open-source projects, Version control systems, Debugging and troubleshooting complex systems
- Demonstrated ability to learn and adapt to new technologies quickly
- Basic understanding of software development lifecycle (SDLC)
- Strong problem-solving and analytical skills
- Excellent written and verbal communication skills
"""

user_prompt=f"""
Analyze the following job description.
{job_description}
"""

user_message={
    "role":"user",
    "content":user_prompt
}

system_message={
    "role":"system",
    "content":system_prompt
}

messages=[system_message,user_message]

response=client.chat.completions.create(
    model=model,
    messages=messages,
    temperature=0,
    response_format=response_format)

answer=response.choices[0].message.content
raw_json=answer
data_file=json.loads(raw_json)
jd=JobDesc(**data_file)


    
#Step 2:Parsing the Resume

class Resume(BaseModel):
    name:str | None
    email:str | None
    phone:str | None
    total_exp:float | None
    skills:list[str]
    experiences:list[str]
    education:list[str]
    projects:list[str]
    certification:list[str]

Resume_schema=Resume.model_json_schema()

system_prompt_resume=f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {Resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
system_message_resume={
    "role":"system",
    "content":system_prompt_resume
}
#Resume Handling
#For PDF type of files
def read_pdf(file_path):
    reader=PdfReader(file_path)
    text=""
    for page in reader.pages:
        page_text=page.extract_text()
        if page_text:
            text=text+page_text+"\n"
    return text

#For Word Files
def read_docx(file_path):
    document=Document(file_path)
    text=""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text=text+paragraph.text+"\n"



    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text=text+cell.text+"\n"
    return text

def read_resume(file_path):
    if file_path.suffix.lower()==".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower()==".docx":
        return read_docx(file_path)
    else:
        return None


resume_folder = Path("Resumes")

for file_path in resume_folder.iterdir():

    # Skip files that are not PDF or DOCX
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue

    print(f"\nProcessing: {file_path.name}")

    # Read resume
    resume_text = read_resume(file_path)

    # Create user prompt
    user_prompt_resume = f"""
    Parse the following resume:

    {resume_text}
    """

    user_message_resume = {
        "role": "user",
        "content": user_prompt_resume
    }

    messages_resume = [
        system_message_resume,
        user_message_resume
    ]

    # Call Groq API
    response_resume = client.chat.completions.create(
        model=model,
        messages=messages_resume,
        temperature=0,
        response_format=response_format
    )

    # Parse response
    ans_res = response_resume.choices[0].message.content
    data_file_res = json.loads(ans_res)
    resume_op = Resume(**data_file_res)

    # Print result
    print(resume_op)
    print("-" * 80)
